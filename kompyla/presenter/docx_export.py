"""Render a markdown document to .docx using python-docx.

Supports the markdown subset Kompyla actually emits: H1–H4 headings,
paragraphs, bullet lists, and fenced code blocks. Other constructs are
emitted as plain paragraphs.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
_HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")
_FENCE_RE = re.compile(r"^```")


def _strip_inline(text: str) -> str:
    """Strip the basic inline syntax we don't render specially."""
    text = re.sub(r"`([^`]+)`", r"\1", text)        # inline code
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)  # bold
    text = re.sub(r"\*([^*]+)\*", r"\1", text)      # italic
    text = re.sub(r"\[\[([^\]]+)\]\]", r"\1", text) # wiki links
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # md links
    return text


def md_to_docx(markdown_text: str, out_path: Path, *, title: str | None = None) -> Path:
    doc = Document()

    if title:
        doc.add_heading(title, level=0)

    in_code = False
    code_lines: list[str] = []

    for raw_line in markdown_text.splitlines():
        # Skip YAML frontmatter
        if raw_line.startswith("---") and not in_code:
            continue

        if _FENCE_RE.match(raw_line):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(raw_line)
            continue

        line = raw_line.rstrip()

        if not line.strip():
            continue

        if m := _HEADING_RE.match(line):
            level = len(m.group(1))
            doc.add_heading(_strip_inline(m.group(2)), level=min(level, 4))
            continue

        if m := _BULLET_RE.match(line):
            doc.add_paragraph(_strip_inline(m.group(1)), style="List Bullet")
            continue

        if m := _NUMBERED_RE.match(line):
            doc.add_paragraph(_strip_inline(m.group(1)), style="List Number")
            continue

        doc.add_paragraph(_strip_inline(line))

    if in_code and code_lines:
        p = doc.add_paragraph()
        p.add_run("\n".join(code_lines)).font.name = "Courier New"

    doc.save(out_path)
    return out_path
